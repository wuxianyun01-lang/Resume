# -*- coding: utf-8 -*-
"""Generate Wu Zhiyong's IT operations resumes as DOCX.

Design tokens follow the documents skill's compact_reference_guide preset with
a named resume override: A4 portrait, Microsoft YaHei, cyan section headings,
single-column body, and a one-row header grid only for the photo placeholder.
"""

import os

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


INK = RGBColor(0x10, 0x18, 0x20)
MUTED = RGBColor(0x5C, 0x6B, 0x76)
TEAL = RGBColor(0x0E, 0x74, 0x90)
TEAL_DARK = RGBColor(0x15, 0x5E, 0x75)
HEADER_TEAL = RGBColor(0x11, 0x18, 0x20)
AMBER = RGBColor(0xD9, 0x77, 0x06)
KEYWORD_GRAY = RGBColor(0x4A, 0x5A, 0x63)

FONT = "微软雅黑"


def cm_to_dxa(cm):
    return int(round(cm * 566.929133858))


def set_run_font(run, name=FONT, size=10, color=INK, bold=False, italic=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:eastAsia"), name)


def set_style_font(style, name=FONT, size=10, color=INK, bold=False):
    style.font.name = name
    style.font.size = Pt(size)
    style.font.color.rgb = color
    style.font.bold = bold
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:eastAsia"), name)


def paragraph_border_bottom(paragraph, color="B8DCE6", size="8", space="4"):
    ppr = paragraph._p.get_or_add_pPr()
    pbdr = ppr.find(qn("w:pBdr"))
    if pbdr is None:
        pbdr = OxmlElement("w:pBdr")
        ppr.append(pbdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)
    pbdr.append(bottom)


def add_paragraph(doc, text, size=10, color=INK, bold=False, before=0, after=3,
                  line=1.15, style=None, align=None):
    p = doc.add_paragraph(style=style)
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    if align is not None:
        pf.alignment = align
    run = p.add_run(text)
    set_run_font(run, size=size, color=color, bold=bold)
    return p


def add_rich_paragraph(doc, parts, before=0, after=3, line=1.15, style=None):
    p = doc.add_paragraph(style=style)
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    for text, kwargs in parts:
        run = p.add_run(text)
        set_run_font(run, **kwargs)
    return p


def add_section(doc, title, before=8, size=13):
    p = add_paragraph(doc, title, size=size, color=TEAL_DARK, bold=True, before=before, after=3, line=1.0)
    paragraph_border_bottom(p)
    return p


def add_bullets(doc, items, size=10, after=3, line=1.18):
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(after)
        pf.line_spacing = line
        pf.left_indent = Cm(0.55)
        pf.first_line_indent = Cm(-0.25)
        for run in p.runs:
            set_run_font(run, size=size)


def add_entry_header(doc, title, meta, size=10.5):
    add_rich_paragraph(
        doc,
        [
            (title, {"size": size, "color": INK, "bold": True}),
            ("　" + meta, {"size": 9.5, "color": MUTED}),
        ],
        before=2,
        after=2,
        line=1.1,
    )


def add_photo_placeholder_table(doc, usable_width_cm, photo_height_cm=3.6):
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.style = "Table Grid"
    table.alignment = WD_ALIGN_PARAGRAPH.LEFT
    left_w = usable_width_cm - 3.2
    right_w = 3.2
    table.columns[0].width = Cm(left_w)
    table.columns[1].width = Cm(right_w)

    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(cm_to_dxa(usable_width_cm)))
    tbl_w.set(qn("w:type"), "dxa")

    grid = tbl.find(qn("w:tblGrid"))
    if grid is not None:
        cols = grid.findall(qn("w:gridCol"))
        if len(cols) == 2:
            cols[0].set(qn("w:w"), str(cm_to_dxa(left_w)))
            cols[1].set(qn("w:w"), str(cm_to_dxa(right_w)))

    left_cell, right_cell = table.rows[0].cells
    left_cell.width = Cm(left_w)
    right_cell.width = Cm(right_w)

    left_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    right_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    left_para = left_cell.paragraphs[0]
    left_para.paragraph_format.space_before = Pt(0)
    left_para.paragraph_format.space_after = Pt(0)
    left_para.paragraph_format.line_spacing = 1.12
    left_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    add_header_text_to_cell(left_cell)

    right_para = right_cell.paragraphs[0]
    right_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    right_para.paragraph_format.space_before = Pt(0)
    right_para.paragraph_format.space_after = Pt(0)
    run = right_para.add_run("照片")
    set_run_font(run, size=11, color=MUTED, bold=True)
    hint = right_cell.add_paragraph("3.0 × 4.0 cm")
    hint.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hint.paragraph_format.space_before = Pt(4)
    hint.paragraph_format.space_after = Pt(0)
    set_run_font(hint.runs[0], size=8, color=MUTED)

    row = table.rows[0]
    row.height = Cm(photo_height_cm)
    row.height_rule = 1  # atLeast
    return table


def add_header_text_to_cell(cell):
    name_para = cell.paragraphs[0]
    run = name_para.add_run("吴志勇")
    set_run_font(run, size=22, color=HEADER_TEAL, bold=True)

    role_para = cell.add_paragraph()
    role_para.paragraph_format.space_before = Pt(2)
    role_para.paragraph_format.space_after = Pt(6)
    role_para.paragraph_format.line_spacing = 1.1
    run = role_para.add_run("IT运维工程师 · 桌面运维 / 自动化 / AI 应用")
    set_run_font(run, size=11, color=TEAL, bold=True)

    meta_para = cell.add_paragraph()
    meta_para.paragraph_format.space_after = Pt(0)
    meta_para.paragraph_format.line_spacing = 1.35
    run = meta_para.add_run("电话：19170245568　邮箱：2268562600@qq.com")
    set_run_font(run, size=8.5, color=MUTED)

    meta2 = cell.add_paragraph()
    meta2.paragraph_format.space_after = Pt(0)
    meta2.paragraph_format.line_spacing = 1.35
    run = meta2.add_run("微信：19170245568　城市：福建宁德")
    set_run_font(run, size=8.5, color=MUTED)

    meta3 = cell.add_paragraph()
    meta3.paragraph_format.space_after = Pt(0)
    meta3.paragraph_format.line_spacing = 1.35
    run = meta3.add_run("政治面貌：中共党员　籍贯：江西九江　出生年月：2002.08")
    set_run_font(run, size=8.5, color=MUTED)


def setup_document(margins_cm):
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(margins_cm)
    section.bottom_margin = Cm(margins_cm)
    section.left_margin = Cm(margins_cm + 0.3)
    section.right_margin = Cm(margins_cm + 0.3)
    section.header_distance = Cm(0.8)
    section.footer_distance = Cm(0.8)

    normal = doc.styles["Normal"]
    set_style_font(normal, size=10, color=INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(3)
    normal.paragraph_format.line_spacing = 1.15

    doc.core_properties.title = "吴志勇 - IT运维工程师简历"
    doc.core_properties.author = "吴志勇"
    doc.core_properties.subject = "桌面运维 / 自动化 / AI 应用"
    return doc


def add_keyword_block(doc, after_spacing=0, compact=False):
    text = (
        "【岗位关键词】IT运维、桌面运维、Windows终端运维、装机自动化、批量部署、系统异常排查、"
        "Python自动化、飞书API、内部网页、资产全生命周期管理、AI大模型接入、SOP标准化、知识库"
        if compact
        else "【岗位关键词】IT运维、桌面运维、Windows终端运维、装机自动化、软件批量部署、系统异常排查、"
        "办公外设运维、Python自动化、飞书机器人API、内部网页、HTML页面、IT资产全生命周期管理、"
        "AI大模型接入、SOP标准化、知识库、会议室网络运维、批量配置部署"
    )
    p = add_paragraph(
        doc,
        text,
        size=7 if compact else 8,
        color=KEYWORD_GRAY,
        before=2,
        after=after_spacing,
        line=1.1 if compact else 1.15,
    )
    return p


def build_one_page(path):
    doc = setup_document(margins_cm=0.4)
    usable_width_cm = 21.0 - 2 * (0.4 + 0.3)
    add_photo_placeholder_table(doc, usable_width_cm, photo_height_cm=2.8)

    add_section(doc, "个人定位", before=2, size=12)
    add_paragraph(
        doc,
        "一年宁德新能源 IT 运维经验 + 半年硬件测试背景。桌面运维扎实，擅长 Python 自动化与轻量网页开发，"
        "熟悉 AI 大模型接入；习惯把零散运维场景 SOP 化、标准化、数字化。",
        size=9,
        after=0.5,
        line=1.1,
    )

    add_section(doc, "核心成果", before=3, size=12)
    add_bullets(
        doc,
        [
            "一周独立完成 200+ 台校招 PC 标准化配置交付，零设备故障、零兼容问题。",
            "自研打印机智能监控系统，设备故障响应时长缩短 90% 以上。",
            "主导 IT 仓库数字化改造，资产盘点效率提升 70% 以上。",
            "自研苹果文件转换等工具，跨系统文件转换效率提升 80%。",
        ],
        size=9,
        after=0.5,
        line=1.05,
    )

    add_section(doc, "工作经历", before=3, size=12)
    add_entry_header(doc, "桌面运维资深技术员", "宁德新能源科技有限公司 · IDT 部门 · 2025.07 - 至今", size=10)
    add_bullets(
        doc,
        [
            "Windows 终端运维、装机自动化与批量部署，负责 IT 仓库设备配置与退仓检测，日均出仓约 40 台；SecureCRT 完成会议室 VLAN 规划。",
            "Python 自研打印机监控、文件快搜、苹果转换、域账号一键改密等工具，对接飞书机器人自动告警。",
            "主导 IT 仓库数字化、资产全生命周期、内网导航页面与 AI 大模型接入，沉淀全套 SOP 与知识库。",
        ],
        size=9,
        after=0.5,
        line=1.05,
    )
    add_paragraph(
        doc,
        "2025.03 - 2025.07　湖北合力源电子技术有限公司 · 硬件测试工程师：驻点武汉理工大学，完成新车型胎压测试项目，"
        "负责线束制作、程序烧录、路测数据采集与 MATLAB 数据分析。",
        size=9,
        after=0.5,
        line=1.05,
    )

    add_section(doc, "项目经历", before=3, size=12)
    add_entry_header(doc, "打印机智能监控自动化系统", "Python / 飞书 API", size=10)
    add_bullets(
        doc,
        ["数据库数据抓取与解析，飞书机器人实时告警，故障响应时长缩短 90% 以上。"],
        size=9,
        after=0.5,
        line=1.05,
    )
    add_entry_header(doc, "装机自动化与办公工具集", "Python / 自动化脚本", size=10)
    add_bullets(
        doc,
        ["装机自动化脚本支撑日均出仓约 40 台；自研文件快搜、苹果文件转换、域账号一键改密。"],
        size=9,
        after=0.5,
        line=1.05,
    )
    add_entry_header(doc, "IT 仓库数字化升级与资产全生命周期管理", "资产台账 / SOP", size=10)
    add_bullets(
        doc,
        ["全流程 SOP + 数字化台账 + 可视化导航，资产盘点效率提升 70% 以上。"],
        size=9,
        after=0.5,
        line=1.05,
    )

    add_section(doc, "专业技能与证书", before=3, size=12)
    add_bullets(
        doc,
        [
            "桌面运维与网络：Windows 终端运维、批量部署、会议室 VLAN、退仓检测、办公外设运维。",
            "开发与数据：Python、C、SQL、MATLAB、飞书 API、HTML/CSS、Agent 与工作流。",
            "资产与 AI：资产全生命周期、SOP、知识库、AI 大模型接入；证书：CET-4、高低压电工证、AutoCAD、5G 承载网络运维（中级）。",
        ],
        size=9,
        after=0.5,
        line=1.05,
    )

    add_section(doc, "其他经历、荣誉与教育", before=3, size=12)
    add_paragraph(
        doc,
        "2023.06 - 2023.09　猿起（武汉）科技有限公司 · 班主任兼销售：每周服务 60+ 客户，多期销售榜前列，月均 GMV 约 2 万。",
        size=9,
        after=0.5,
        line=1.05,
    )
    add_paragraph(
        doc,
        "2023.09 - 2025.06　班长兼学习委员；2021.10 - 2022.11　青年志愿者协会会长。",
        size=9,
        after=0.5,
        line=1.05,
    )
    add_paragraph(
        doc,
        "2021.09 - 2025.06　武汉东湖学院 · 通信工程 · 本科；成绩专业第一、连续两年国家励志奖学金；"
        "大唐杯全国一等奖、数学建模省级一等奖等。",
        size=9,
        after=0.5,
        line=1.05,
    )

    add_keyword_block(doc, compact=True)
    doc.save(path)


def build_two_page(path):
    doc = setup_document(margins_cm=0.8)
    usable_width_cm = 21.0 - 2 * (0.8 + 0.3)
    add_photo_placeholder_table(doc, usable_width_cm, photo_height_cm=3.3)

    add_section(doc, "个人定位", before=2, size=12)
    add_paragraph(
        doc,
        "一年宁德新能源 IDT 部门 IT 运维经验，叠加半年硬件测试背景，兼具桌面运维落地能力、Python 自动化开发能力"
        "与 AI 应用接入能力。擅长从业务痛点出发，把零散运维工作、高频办公场景、故障处置流程 SOP 化、标准化、体系化，"
        "并独立落地打印机智能监控、装机自动化、苹果文件转换、IT 仓库数字化等内部自研项目，以标准化流程 + 技术自研"
        "+ 智能化升级提升企业 IT 服务效率。",
        size=9.5,
        after=2,
        line=1.18,
    )

    add_section(doc, "工作经历", before=5, size=12)
    add_entry_header(doc, "桌面运维资深技术员", "宁德新能源科技有限公司 · IDT 部门 · 2025.07 - 至今")
    add_bullets(
        doc,
        [
            "负责公司全域 Windows 终端运维、软件批量部署与系统异常排查；负责 IT 仓库电脑/终端机安装配置与退仓检测，日均出仓约 40 台，基于服务器编写装机自动化脚本实现装机自动化。",
            "基于 Python 独立开发打印机智能监控系统并对接飞书机器人实现自动告警与 7×24 小时无人值守监控；自研文件快搜、苹果文件转换、域账号一键修改密码等办公工具。",
            "主导 IT 仓库数字化改造与资产全生命周期管理，落地自助导航与实景可视化方案，盘点效率提升 70% 以上。",
            "使用 SecureCRT 命令行修改会议室 VLAN，完成厂区会议室规划与网络运维；沉淀全套运维 SOP 与问题知识库。",
            "独立承接 200+ 台校招 PC 批量配置项目，一周完成系统安装、软件部署、驱动适配与资产绑定，零故障交付。",
            "基于公司内网制作导航类 HTML 页面；完成 AI 大模型 API 接入、客户端部署、密钥与权限调试及异常排错。",
        ],
        size=9.5,
        after=1.5,
        line=1.14,
    )

    add_entry_header(doc, "硬件测试工程师", "湖北合力源电子技术有限公司 · 2025.03 - 2025.07")
    add_bullets(
        doc,
        [
            "驻点武汉理工大学，对接浙江万安科技新车型胎压测试项目，跟进硬件制作、软件集成、路测数据采集全流程。",
            "完成线束制作、程序烧录与不同胎压工况数据采集，使用 MATLAB 分析数据，配合项目验收。",
        ],
        size=9.5,
        after=1.5,
        line=1.14,
    )

    add_section(doc, "项目经历", before=5, size=12)
    projects = [
        (
            "打印机智能监控自动化系统",
            "Python / 飞书 API",
            "针对人工巡检效率低、故障发现滞后等问题，独立开发全套智能监控系统，替代人工巡检。",
            [
                "编写数据抓取与解析逻辑，自动读取数据库内打印机运行数据，精准识别离线、卡纸、缺墨、故障停机等异常。",
                "完成飞书机器人 API 深度对接，实现 7×24 小时无人值守监控与实时告警，故障响应时长缩短 90% 以上。",
            ],
        ),
        (
            "苹果文件格式转换工具",
            "Python",
            "针对 Mac 与 Windows 文件格式不兼容、转换繁琐、第三方工具不安全等问题，自研轻量内部转换系统。",
            [
                "基于 Python 实现多类型文件一键批量转换，适配日常办公文件使用场景。",
                "本地安全转换，规避第三方工具数据泄露风险，员工文件转换效率提升 80%。",
            ],
        ),
        (
            "装机自动化脚本与办公工具集",
            "Python / 自动化脚本",
            "基于公司服务器沉淀批量装机与高频办公自动化能力，降低重复人工操作。",
            [
                "编写装机自动化脚本，实现电脑装机、环境配置流程自动化，支撑日均出仓约 40 台。",
                "Python 自研文件快搜、苹果文件转换、域账号一键修改密码等工具，显著提升办公效率。",
            ],
        ),
        (
            "IT 仓库数字化升级与资产全生命周期管理",
            "资产台账 / SOP / 可视化导航",
            "针对 IT 仓库资产杂乱、查找困难、台账不规范等问题，重构标准化、数字化资产管控模式。",
            [
                "梳理落地入库、领用、调拨、盘点、报废全流程 SOP；规划设计自助导航与实景可视化导航方案。",
                "数字化资产台账实现实物与数据一一对应，盘点效率提升 70% 以上，杜绝资产流失与闲置。",
            ],
        ),
        (
            "企业 AI 大模型能力接入与落地",
            "API 接入 / 部署调试",
            "为赋能企业智能化办公，落地内部 AI 应用能力，完成大模型 API 对接、环境部署与异常运维。",
            [
                "完成多厂商 AI 大模型客户端部署、API 密钥配置、权限调试与接口对接。",
                "排查模型连接异常与参数适配问题，搭建稳定的企业内部 AI 使用环境，支撑办公场景落地。",
            ],
        ),
        (
            "校招 200+ 台 PC 标准化配置部署",
            "批量部署 / 资产管理",
            "针对校招新人大批量办公设备需求，快速完成数百台 PC 的系统安装、软件部署、环境调试与资产绑定。",
            [
                "制定统一配置规范，单人一周完成 200+ 台设备系统装机、软件批量部署、驱动适配与资产台账绑定。",
                "零设备故障、零兼容问题，保障新人入职即用，沉淀可复用的 PC 批量配置标准化流程。",
            ],
        ),
    ]
    for title, meta, desc, bullets in projects:
        add_entry_header(doc, title, meta)
        add_paragraph(doc, desc, size=9, after=1, line=1.12)
        add_bullets(doc, bullets, size=9, after=1, line=1.12)

    add_section(doc, "专业技能与证书", before=5, size=12)
    skills = [
        "开发能力：熟练 Python、C 语言、SQL、MATLAB、Proteus8、STM32、AutoCAD，可独立完成办公工具与监控系统自研。",
        "自动化与智能化：飞书机器人 API、自动化告警、装机自动化脚本、Agent 与工作流；AI 大模型 API 接入、部署与调参。",
        "运维能力：Windows 终端运维、软件批量部署、会议室网络运维（VLAN）、设备退仓检测与办公外设运维。",
        "资产管理：IT 资产全生命周期管理、仓库数字化优化、自助导航与可视化导航方案落地。",
        "综合能力：大厂标准化工作思维，擅长 SOP 标准化、知识库搭建、需求对接与项目全流程落地。",
        "证书：CET-4、低压/高压电工证、AutoCAD 工程师证、5G 承载网络运维职业技能等级证书（中级）。",
    ]
    add_bullets(doc, skills, size=9, after=1, line=1.12)

    add_section(doc, "校园经历", before=5, size=12)
    add_paragraph(
        doc,
        "2023.09 - 2025.06　班长兼学习委员：组织班级事务，带动学习氛围，带领同学参加学科竞赛；"
        "2021.10 - 2022.11　青年志愿者协会会长：带领 10+ 人团队组织防疫、地铁、社区志愿服务，"
        "参与创新创业项目并获学校创业扶持。",
        size=9,
        after=2,
        line=1.12,
    )

    add_section(doc, "获奖与荣誉", before=5, size=12)
    honors = [
        "2024 全国大学生文学作品大赛全国三等奖；2023 新华三杯全国大学生数字技术大赛省级二等奖。",
        "2022 第九届“大唐杯”全国大学生移动通信 5G 技术大赛全国一等奖；2021 全国大学生数学建模大赛省级一等奖。",
        "连续两年国家励志奖学金、校奖学金；优秀毕业生、优秀毕业论文、优秀志愿者、优秀团干、优秀班干、三好学生、学习之星；成绩专业第一。",
    ]
    add_bullets(doc, honors, size=9, after=1, line=1.12)

    add_section(doc, "其他经历", before=5, size=12)
    add_paragraph(
        doc,
        "2023.06 - 2023.09　猿起（武汉）科技有限公司 · 班主任兼销售：每周服务 60+ 客户，多期销售榜前列，"
        "月均 GMV 约 2 万，最高 3.5 万。",
        size=9,
        after=2,
        line=1.12,
    )

    add_section(doc, "教育背景", before=5, size=12)
    add_paragraph(
        doc,
        "2021.09 - 2025.06　武汉东湖学院 · 通信工程 · 本科 · 中共党员",
        size=9.5,
        after=1,
        line=1.12,
    )
    add_paragraph(
        doc,
        "主修课程：Python 程序设计、单片机 C 语言程序设计、嵌入式系统、移动通信、电子技术、PCB 设计、工程制图与 CAD。",
        size=9,
        after=2,
        line=1.12,
    )

    add_section(doc, "自我评价", before=5, size=12)
    add_paragraph(
        doc,
        "1.5 年 IT 运维与硬件测试实战经验，兼具技术开发能力、大批量设备运维落地能力与标准化体系化工作思维。"
        "擅长 Python 自动化、装机自动化、PC 批量部署、办公场景数字化改造与 SOP 沉淀；在校成绩专业第一，"
        "连续两年国家励志奖学金，获大唐杯全国一等奖、数学建模省级一等奖。善于把标准化流程 + 技术自研 + 智能化升级"
        "用于解决实际运维痛点，能够高效推进企业信息化、标准化、数字化建设。",
        size=9,
        after=2,
        line=1.12,
    )

    add_keyword_block(doc)
    doc.save(path)


def main():
    root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    word_dir = os.path.join(root, "word")
    os.makedirs(word_dir, exist_ok=True)
    build_one_page(os.path.join(word_dir, "吴志勇-IT运维工程师-一页版.docx"))
    build_two_page(os.path.join(word_dir, "吴志勇-IT运维工程师-两页版.docx"))
    print("Generated both DOCX resumes.")


if __name__ == "__main__":
    main()
