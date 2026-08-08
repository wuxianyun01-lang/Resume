# -*- coding: utf-8 -*-
"""Automated QA for the resume project (web + Word + PDF)."""

import os
import re

import pdfplumber
from docx import Document


ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
WORD_DIR = os.path.join(ROOT, "word")


def docx_text(path):
    doc = Document(path)
    parts = []
    for p in doc.paragraphs:
        if p.text:
            parts.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if p.text:
                        parts.append(p.text)
    return "\n".join(parts)


def html_text(path):
    with open(path, encoding="utf-8") as f:
        html = f.read()
    text = re.sub(r"<script.*?</script>", " ", html, flags=re.S)
    text = re.sub(r"<style.*?</style>", " ", text, flags=re.S)
    text = re.sub(r"<[^>]+>", " ", text)
    text = re.sub(r"\s+", " ", text)
    return text


KEY_FACTS = [
    "吴志勇",
    "19170245568",
    "2268562600@qq.com",
    "福建宁德",
    "武汉东湖学院",
    "通信工程",
    "2021.09 - 2025.06",
    "2025.07",
    "90%",
    "70%",
    "80%",
    "200+",
    "猿起",
    "打印机智能监控",
    "苹果文件",
    "仓库数字化",
    "AI 大模型",
    "飞书",
    "导航",
    "校招",
    "SOP",
    "湖北合力源",
    "硬件测试",
    "装机自动化",
    "域账号",
    "VLAN",
    "大唐杯",
    "数学建模",
    "国家励志奖学金",
    "中共党员",
    "班长",
    "青年志愿者协会",
    "CET-4",
    "5G 承载网络运维",
]


def check_facts(label, text):
    missing = [fact for fact in KEY_FACTS if fact not in text]
    if missing:
        raise AssertionError(f"{label} missing facts: {missing}")
    if "网易" in text:
        raise AssertionError(f"{label} still contains 网易")


def check_pdf_bounds(pdf_path, left_pt, right_pt, top_pt, bottom_pt):
    with pdfplumber.open(pdf_path) as pdf:
        for i, page in enumerate(pdf.pages):
            words = page.extract_words()
            for word in words:
                if word["x0"] < left_pt - 2 or word["x1"] > page.width - right_pt + 2:
                    raise AssertionError(
                        f"{os.path.basename(pdf_path)} page {i + 1} word outside horizontal margin: {word['text']}"
                    )
                if word["top"] < top_pt - 2 or word["bottom"] > page.height - bottom_pt + 2:
                    raise AssertionError(
                        f"{os.path.basename(pdf_path)} page {i + 1} word outside vertical margin: {word['text']}"
                    )


def main():
    one_docx = os.path.join(WORD_DIR, "吴志勇-IT运维工程师-一页版.docx")
    two_docx = os.path.join(WORD_DIR, "吴志勇-IT运维工程师-两页版.docx")
    one_pdf = os.path.join(WORD_DIR, "吴志勇-IT运维工程师-一页版.pdf")
    two_pdf = os.path.join(WORD_DIR, "吴志勇-IT运维工程师-两页版.pdf")
    index_html = os.path.join(ROOT, "index.html")

    one_text = docx_text(one_docx)
    two_text = docx_text(two_docx)
    web_text = html_text(index_html)

    check_facts("web", web_text)
    check_facts("one-page docx", one_text)
    check_facts("two-page docx", two_text)

    with pdfplumber.open(one_pdf) as pdf:
        assert len(pdf.pages) == 1, "one-page PDF must be exactly 1 page"
    with pdfplumber.open(two_pdf) as pdf:
        assert len(pdf.pages) == 2, "two-page PDF must be exactly 2 pages"

    # One-page: top/bottom 0.4cm, left/right 0.7cm.
    check_pdf_bounds(one_pdf, left_pt=19.84, right_pt=19.84, top_pt=11.34, bottom_pt=11.34)
    # Two-page: top/bottom 0.8cm, left/right 1.1cm.
    check_pdf_bounds(two_pdf, left_pt=31.18, right_pt=31.18, top_pt=22.68, bottom_pt=22.68)

    print("All resume QA checks passed.")


if __name__ == "__main__":
    main()
